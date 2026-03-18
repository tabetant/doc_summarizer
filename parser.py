import re
from pathlib import Path


class ParserError(Exception):
    pass


def parse_document(file_path: str) -> str:
    path = Path(file_path)

    if not path.exists():
        raise ParserError(f"File not found: {file_path}")

    ext = path.suffix.lower()
    if ext == ".pdf":
        text = _parse_pdf(path)
    elif ext in (".docx", ".doc"):
        text = _parse_docx(path)
    elif ext == ".txt":
        text = _parse_txt(path)
    else:
        raise ParserError(
            f"Unsupported file type '{ext}'. Supported types: .pdf, .docx, .doc, .txt"
        )

    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    if len(text) < 50:
        raise ParserError(
            "Document appears to be empty or unreadable — possibly a scanned image PDF."
        )

    return text


def _parse_pdf(path: Path) -> str:
    import pypdf

    try:
        reader = pypdf.PdfReader(str(path))
    except pypdf.errors.FileNotDecryptedError:
        raise ParserError("PDF is password-protected and cannot be read.")

    pages = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            pages.append(page_text)

    return "\n\n".join(pages)


def _parse_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))

    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_parts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_parts:
                parts.append(" | ".join(row_parts))

    return "\n".join(parts)


def _parse_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")
